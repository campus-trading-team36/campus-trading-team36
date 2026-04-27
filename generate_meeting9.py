from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('COMP208 Project \u2014 Team Minutes')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Courier New'

doc.add_paragraph()

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Meeting 9 Minutes')
run2.font.size = Pt(16)
run2.font.name = 'Courier New'

doc.add_paragraph()

# Info table
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'

info = [
    ('Module / Project', 'COMP208 Project'),
    ('Team number', '36'),
    ('Meeting number', '9'),
    ('Date', '24 April 2026'),
    ('Time', '10:30\u201312:00'),
    ('Location', 'Online'),
    ('Chair', 'Pengyu Chen'),
    ('Minute-talker', 'Pengyu Chen'),
    ('Present', 'Futian Bi, Pengyu Chen'),
]

for i, (label, value) in enumerate(info):
    row = table.rows[i]
    # Label cell
    lc = row.cells[0]
    lc.text = label
    lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lc.paragraphs[0].runs[0]
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    set_cell_border(lc)

    # Value cell
    vc = row.cells[1]
    vc.text = value
    run2 = vc.paragraphs[0].runs[0]
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)
    set_cell_border(vc)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(9)

doc.add_paragraph()

# Agenda
def add_heading(doc, text, bold=True, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    return p

add_heading(doc, 'Agenda (shared before the meeting)')

agenda_items = [
    'Actions from previous meeting',
    'Summary of current work progress',
    'Review of frontend completion and improvements',
    'Back-end development update',
    'System testing results',
    'Task allocation for final phase',
    'Next meeting',
]
for item in agenda_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

# Horizontal rule (simulate with a border paragraph)
doc.add_paragraph()
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:color'), '000000')
pBdr.append(bottom)
pPr.append(pBdr)

# Minutes heading
add_heading(doc, 'Minutes')

doc.add_paragraph()

# Section 1
add_heading(doc, '1. Actions from previous meeting (Meeting Eight)')
p = doc.add_paragraph()
run = p.add_run(
    'Actions: Futian Bi to build the HTML structure of all pages, set up basic '
    'layout and organise page content \u2014 Completed.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run(
    'Actions: Pengyu Chen to implement CSS styling, add JavaScript interactions, '
    'optimise visual effects and ensure stable display across pages \u2014 Completed.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run(
    'Actions: Both members to cooperate on frontend testing, bug fixing and '
    'performance optimisation \u2014 Completed.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 2
add_heading(doc, '2. Summary of current work progress')
p = doc.add_paragraph()
run = p.add_run(
    'Each member reported their individual progress. Futian Bi confirmed that '
    'all HTML page structures have been completed, including the homepage, product '
    'listing page, product detail page, shopping cart page and user profile page. '
    'Pengyu Chen confirmed that CSS styling and JavaScript interactions have been '
    'fully applied across all pages, with responsive design verified on multiple '
    'screen sizes. Both members agreed that the frontend development phase is '
    'largely complete and ready for integration with the backend.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 3
add_heading(doc, '3. Review of frontend completion and improvements')
p = doc.add_paragraph()
run = p.add_run(
    'The team reviewed the completed frontend pages together. Several minor UI '
    'inconsistencies were identified and corrected, including button alignment on '
    'the cart page and font sizing in the navigation bar. The team also agreed to '
    'improve the loading performance of product images and to add basic form '
    'validation feedback messages for the login and registration pages. These '
    'improvements were assigned and are expected to be completed before the next '
    'meeting.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 4
add_heading(doc, '4. Back-end development update')
p = doc.add_paragraph()
run = p.add_run(
    'Pengyu Chen provided an update on the current state of the backend. The '
    'Node.js server and database connection have been set up. Core API endpoints '
    'for user authentication, product browsing, and shopping cart management are '
    'functional. Integration between the frontend pages and the backend APIs is '
    'ongoing. Remaining backend tasks include the seller rating system, order '
    'history retrieval, and admin management functions.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 5
add_heading(doc, '5. System testing results')
p = doc.add_paragraph()
run = p.add_run(
    'Initial system testing was conducted covering user registration, login, '
    'product listing, and cart operations. The majority of tested functions '
    'performed as expected. Two bugs were identified: an intermittent session '
    'timeout issue during checkout and an image upload failure for listings with '
    'more than three images. Both issues have been logged and will be prioritised '
    'for fixing before the final submission. Further end-to-end testing will '
    'continue in the next development cycle.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 6
add_heading(doc, '6. Task allocation for final phase')
p = doc.add_paragraph()
run = p.add_run(
    'Futian Bi is responsible for fixing the identified frontend UI issues, '
    'improving image loading performance, and assisting with end-to-end testing. '
    'Pengyu Chen is responsible for resolving the two reported backend bugs, '
    'completing the remaining API endpoints, and finalising backend-frontend '
    'integration. Both members will jointly conduct comprehensive system testing '
    'and prepare project documentation for final submission.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

# Section 7 - AOB
p = doc.add_paragraph()
run = p.add_run('7. Any other business')
run.bold = True
run.font.name = 'Courier New'
run.font.size = Pt(10)
run2 = p.add_run('  No additional issues were raised during this meeting.')
run2.font.name = 'Courier New'
run2.font.size = Pt(10)

doc.add_paragraph()

# Section 8 - Next meeting
add_heading(doc, '8. Next meeting')
p = doc.add_paragraph()
run = p.add_run(
    'The next meeting will verify that all identified bugs have been resolved, '
    'conduct a final round of system testing, and review the completed '
    'documentation. The team will also confirm the final submission checklist '
    'and ensure all deliverables are in order.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Bottom rule
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:color'), '000000')
pBdr.append(bottom)
pPr.append(pBdr)

output_path = r'E:\campus-trading-team36\Team36-Minutes-20260424.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
